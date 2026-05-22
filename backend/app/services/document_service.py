import io
import os

from pypdf import PdfReader


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    text = ""

    if ext == ".pdf":
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    elif ext == ".docx":
        from docx import Document as DocxDocument
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract text from tables inside the Word document
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
    elif ext == ".xlsx":
        from openpyxl import load_workbook
        xlsx_file = io.BytesIO(file_content)
        wb = load_workbook(xlsx_file, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text += f"--- Sheet: {sheet_name} ---\n"
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_vals):
                    text += " | ".join(row_vals) + "\n"
        wb.close()
    elif ext in [".txt", ".csv"]:
        text = file_content.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return text


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200):
    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap

    return chunks
